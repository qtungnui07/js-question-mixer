import json
import random
import re
import sys
import os
from copy import deepcopy
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def load_questions(path):
    with open(path, "r", encoding="utf-8") as f:
        data = f.read()
    if data.strip().startswith("const"):
        data = data.split("=", 1)[1].strip()
        if data.endswith(";"):
            data = data[:-1]
    return json.loads(data)
def clean_answer(ans):
    return re.sub(r"^[A-D][\.\)]?\s*", "", ans).strip()
def shuffle_answers(answers):
    texts = [clean_answer(a) for a in answers]
    random.shuffle(texts)
    letters = ["A", "B", "C", "D"]
    return [f"{letters[i]}. {texts[i]}" for i in range(len(texts))]
def generate_versions(questions, num_versions):
    versions = []
    for _ in range(num_versions):
        new_questions = []
        for q in questions:
            nq = deepcopy(q)
            nq["answers"] = shuffle_answers(nq["answers"])
            new_questions.append(nq)
        random.shuffle(new_questions)
        versions.append(new_questions)
    return versions
def export_word(versions, output_name):
    os.makedirs("output", exist_ok=True)
    base = output_name.replace(".docx", "")
    for i, questions in enumerate(versions):
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        title = doc.add_paragraph()
        run=title.add_run(f"Mã đề {i+1}")
        run.bold=True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for idx, q in enumerate(questions, 1):
            doc.add_paragraph(f"Câu {idx}: {q['question']}")
            table = doc.add_table(rows=2, cols=2)
            table.cell(0,0).text = q["answers"][0]
            table.cell(0,1).text = q["answers"][1]
            table.cell(1,0).text = q["answers"][2]
            table.cell(1,1).text = q["answers"][3]
        filename = f"output/{base}_{i+1}.docx"
        doc.save(filename)
        print("Created:", filename)
def main():
    if len(sys.argv) != 4:
        print("Usage: convert <input.js> <output_name> <so_ma_de>")
        return
    input_file = sys.argv[1]
    output_name = sys.argv[2]
    num = int(sys.argv[3])
    questions = load_questions(input_file)
    versions = generate_versions(questions, num)
    export_word(versions, output_name)
if __name__ == "__main__":
    main()